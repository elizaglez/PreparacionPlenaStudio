from __future__ import annotations

import json
from pathlib import Path

from app.models import Project
from app.persistence.generated_article_repository import (
    GeneratedArticleRepositoryError,
    JsonGeneratedArticleRepository,
)


class WordExportError(RuntimeError):
    pass


def export_master_to_docx(project: Project) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise WordExportError(
            "No está instalado python-docx. Ejecuta pip install -r requirements.txt."
        ) from exc

    root = Path(project.root)
    master_path = root / "trabajo" / "master.json"
    if not master_path.is_file():
        raise WordExportError("Primero genera el MASTER.")

    try:
        master = json.loads(master_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise WordExportError("master.json contiene datos inválidos.") from exc

    output_dir = root / "salidas"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "MASTER.docx"

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)

    document.add_heading(master.get("title", "MASTER"), level=0)
    intro = master.get("introduction", "").strip()
    if intro:
        document.add_paragraph(intro)

    for item in master.get("answers", []):
        question = item.get("question", "")
        document.add_heading(question, level=1)
        document.add_paragraph(item.get("answer", ""))

        optional = [
            ("Explicación bíblica", item.get("scripture_explanation", "")),
            ("Comparación", item.get("comparison", "")),
            ("Aplicación", item.get("application", "")),
            ("Nota de imagen", item.get("image_note", "")),
        ]
        for label, value in optional:
            if str(value).strip():
                paragraph = document.add_paragraph()
                paragraph.add_run(f"{label}: ").bold = True
                paragraph.add_run(str(value).strip())

        scriptures = item.get("scriptures", [])
        if scriptures:
            paragraph = document.add_paragraph()
            paragraph.add_run("Textos bíblicos: ").bold = True
            paragraph.add_run(", ".join(scriptures))

    conclusion = master.get("conclusion", "").strip()
    if conclusion:
        document.add_heading("Conclusión", level=1)
        document.add_paragraph(conclusion)

    document.save(output_path)
    return output_path


def export_generated_article_to_docx(project: Project) -> Path:
    """Export the new GeneratedArticle artifact without creating master.json."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise WordExportError(
            "No está instalado python-docx. Ejecuta pip install -r requirements.txt."
        ) from exc

    try:
        article = JsonGeneratedArticleRepository(project.root).load()
    except GeneratedArticleRepositoryError as exc:
        raise WordExportError(
            "No se pudo leer articulo_generado.json."
        ) from exc

    output_dir = Path(project.root) / "salidas"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "CONTENIDO_GENERADO.docx"
    temporary_output_path = output_path.with_suffix(
        output_path.suffix + ".tmp"
    )

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)

    document.add_heading(article.title or "Contenido generado", level=0)

    for paragraph in article.introduction.paragraphs:
        text = str(paragraph.get("text", "")).strip()
        if text:
            document.add_paragraph(text)

    def add_question(question, *, level: int) -> None:
        document.add_heading(question.question, level=level)
        if question.answer.strip():
            document.add_paragraph(question.answer.strip())
        if question.application.strip():
            paragraph = document.add_paragraph()
            paragraph.add_run("Aplicación: ").bold = True
            paragraph.add_run(question.application.strip())

    for question in article.introduction.questions:
        add_question(question, level=1)

    for section in article.sections:
        document.add_heading(section.subtitle, level=1)

        if section.heygen_transition:
            paragraph = document.add_paragraph()
            paragraph.add_run("Transición: ").bold = True
            paragraph.add_run(section.heygen_transition.strip())

        for question in section.questions:
            add_question(question, level=2)

        if section.section_summary:
            paragraph = document.add_paragraph()
            paragraph.add_run("Resumen de la sección: ").bold = True
            paragraph.add_run(section.section_summary.strip())

        for box in section.boxes:
            document.add_heading(box.title, level=2)
            if box.explanation.strip():
                document.add_paragraph(box.explanation.strip())

    if article.review_questions:
        document.add_heading("Preguntas de repaso", level=1)
        for question in article.review_questions:
            document.add_paragraph(question, style="List Bullet")

    try:
        temporary_output_path.unlink(missing_ok=True)
        document.save(temporary_output_path)
        temporary_output_path.replace(output_path)
    except Exception as exc:
        try:
            temporary_output_path.unlink(missing_ok=True)
        except OSError:
            pass
        raise WordExportError(
            "No se pudo guardar CONTENIDO_GENERADO.docx. "
            "Cierra el documento si está abierto y vuelve a intentarlo."
        ) from exc

    return output_path
