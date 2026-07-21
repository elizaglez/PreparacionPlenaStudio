from __future__ import annotations

from pathlib import Path
from typing import Any

import fitz
from docx import Document


class BibleReadError(RuntimeError):
    pass


def _read_txt(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise BibleReadError("No se pudo detectar la codificación del archivo de citas.")


def _read_docx(path: Path) -> str:
    try:
        document = Document(path)
    except Exception as exc:
        raise BibleReadError("No se pudo abrir el documento DOCX de citas.") from exc

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def _read_pdf(path: Path) -> str:
    try:
        document = fitz.open(path)
    except Exception as exc:
        raise BibleReadError("No se pudo abrir el PDF de citas.") from exc
    try:
        return "\n\n".join(page.get_text("text").strip() for page in document)
    finally:
        document.close()


def read_bible_source(path: str | Path) -> dict[str, Any]:
    source = Path(path)
    if not source.is_file():
        raise BibleReadError("No se encontró el archivo de citas bíblicas.")

    suffix = source.suffix.lower()
    if suffix == ".txt":
        text = _read_txt(source)
    elif suffix == ".docx":
        text = _read_docx(source)
    elif suffix == ".pdf":
        text = _read_pdf(source)
    else:
        raise BibleReadError(
            "El archivo de citas debe ser TXT, DOCX o PDF."
        )

    text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    non_empty_lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not text:
        raise BibleReadError("El archivo de citas está vacío.")

    return {
        "file": str(source.resolve()),
        "format": suffix.lstrip("."),
        "character_count": len(text),
        "line_count": len(non_empty_lines),
        "text": text,
    }
